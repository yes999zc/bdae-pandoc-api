"""
ESA Report MD → DOCX Conversion Service
Pandoc 转换 + python-docx 后处理（仅动态染色）
"""
from flask import Flask, request, jsonify, send_from_directory
import subprocess, tempfile, os, uuid, traceback

app = Flask(__name__)
OUTPUT_DIR     = '/tmp/pandoc_outputs'
REFERENCE_DOCX = '/app/templates/reference.docx'

os.makedirs(OUTPUT_DIR, exist_ok=True)


@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})


@app.route('/convert', methods=['POST'])
def convert():
    try:
        data          = request.get_json(force=True)
        md_content    = data.get('markdown', '')
        report_no     = data.get('report_no', '')
        property_name = data.get('property_name', '')
        filename      = data.get('filename', 'ESA_Report.docx')

        if not md_content:
            return jsonify({'success': False, 'error': 'markdown field is empty'}), 400

        file_id   = str(uuid.uuid4())[:8]
        safe_name = f"{filename.replace('.docx', '')}_{file_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, safe_name)

        # Step 1: Pandoc MD → DOCX
        with tempfile.NamedTemporaryFile(suffix='.md', mode='w',
                                         encoding='utf-8', delete=False) as f:
            f.write(md_content)
            md_path = f.name

        cmd = ['pandoc', md_path, '-o', docx_path, '--from=markdown', '--to=docx']
        if os.path.exists(REFERENCE_DOCX):
            cmd += [f'--reference-doc={REFERENCE_DOCX}']

        result = subprocess.run(cmd, capture_output=True, text=True)
        os.unlink(md_path)

        if result.returncode != 0:
            return jsonify({'success': False,
                            'error': f'Pandoc error: {result.stderr}'}), 500

        # Step 2: 后处理
        postprocess(docx_path)

        return jsonify({
            'success':      True,
            'download_url': f'http://localhost:5050/files/{safe_name}',
            'filename':     safe_name
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e),
                        'trace': traceback.format_exc()}), 500


@app.route('/files/<filename>')
def download(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


def postprocess(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)

    RANK_COLORS = {
        'HIGH': ('C00000', 'FFFFFF'),
        'MED':  ('FF8C00', 'FFFFFF'),
        'LOW':  ('375623', 'FFFFFF'),
    }

    PAGE_WIDTH_DXA = 9639  # A4 左右各2cm边距

    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'),    str(PAGE_WIDTH_DXA))
        tblW.set(qn('w:type'), 'dxa')

        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'left')

        col_count = len(table.columns)
        col_w     = PAGE_WIDTH_DXA // col_count if col_count > 0 else PAGE_WIDTH_DXA
        tblGrid   = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gc in tblGrid.findall(qn('w:gridCol')):
                gc.set(qn('w:w'), str(col_w))

        for ri, row in enumerate(table.rows):
            for cell in row.cells:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'),    str(col_w))
                tcW.set(qn('w:type'), 'dxa')

                text = cell.text.strip().upper()
                if text in RANK_COLORS:
                    _set_cell_color(cell, *RANK_COLORS[text])
                elif ri == 0:
                    _set_cell_color(cell, '000000', 'FFFFFF')

                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    doc.save(docx_path)


def _set_cell_color(cell, bg_hex, fg_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg_hex)
    tcPr.append(shd)

    r, g, b = int(fg_hex[0:2], 16), int(fg_hex[2:4], 16), int(fg_hex[4:6], 16)
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(r, g, b)
            run.bold = True


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5050, debug=False)