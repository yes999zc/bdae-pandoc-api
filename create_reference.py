"""
生成 Pandoc reference.docx 模板
对齐 BDAE ESA 报告样式 - 包含页脚、Logo、保密声明

输出路径：/app/templates/reference.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 品牌色 ────────────────────────────────────────────────────────────────
BRAND_GREEN = RGBColor(0x19, 0x6B, 0x24)   # #196B24 BDAE 绿
DARK_TEXT   = RGBColor(0x26, 0x26, 0x26)   # #262626 正文近黑
FOOTER_GRAY = RGBColor(0x80, 0x80, 0x80)   # #808080 页脚灰色


def make_reference_docx(output_path='templates/reference.docx'):
    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    # ── 页面：A4 ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.top_margin    = Cm(2.3)
        section.bottom_margin = Cm(2.5)  # 增加底部边距给页脚
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        
        # ── 页脚设置 ───────────────────────────────────────────────────────
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 清除默认内容
        for run in footer_para.runs:
            run.clear()
        
        # 添加文件名和保密声明（左下角）
        run1 = footer_para.add_run("{filename} | PRIVILEGED AND CONFIDENTIAL")
        run1.font.size = Pt(8)
        run1.font.color.rgb = FOOTER_GRAY
        run1.font.name = 'Arial'
        
        # 添加页码（右侧）- 使用制表符定位
        footer_para.add_run("\t\t")  # 制表符分隔
        run2 = footer_para.add_run("PAGE ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = BRAND_GREEN
        run2.font.name = 'Arial'
        
        # 添加页码字段
        add_page_number(footer_para, BRAND_GREEN)
        
        run3 = footer_para.add_run(" OF ")
        run3.font.size = Pt(8)
        run3.font.color.rgb = BRAND_GREEN
        run3.font.name = 'Arial'
        
        # 添加总页数字段
        add_num_pages(footer_para, BRAND_GREEN)
        
        # 添加 Logo（右下角）
        # Logo 将在运行时动态添加，这里预留位置
        # 实际 Logo 通过 python-docx 后处理添加

    # ── Normal 正文 ────────────────────────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name      = 'Arial'
    normal.font.size      = Pt(11)
    normal.font.color.rgb = DARK_TEXT
    _set_east_asia_font(normal, '宋体')
    normal.paragraph_format.space_before = Pt(2.5)
    normal.paragraph_format.space_after  = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.15

    # ── Heading 1：品牌绿，18pt，粗体 ──────────────────────────────────────
    h1 = _get_style(doc, 'Heading 1')
    h1.font.name      = 'Arial'
    h1.font.size      = Pt(18)
    h1.font.bold      = True
    h1.font.color.rgb = BRAND_GREEN
    _set_east_asia_font(h1, '黑体')
    h1.paragraph_format.space_before   = Pt(18)
    h1.paragraph_format.space_after    = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # ── Heading 2：品牌绿，16pt，粗体 ──────────────────────────────────────
    h2 = _get_style(doc, 'Heading 2')
    h2.font.name      = 'Arial'
    h2.font.size      = Pt(16)
    h2.font.bold      = True
    h2.font.color.rgb = BRAND_GREEN
    _set_east_asia_font(h2, '黑体')
    h2.paragraph_format.space_before   = Pt(12)
    h2.paragraph_format.space_after    = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # ── Heading 3：品牌绿，14pt，粗体 ──────────────────────────────────────
    h3 = _get_style(doc, 'Heading 3')
    h3.font.name      = 'Arial'
    h3.font.size      = Pt(14)
    h3.font.bold      = True
    h3.font.color.rgb = BRAND_GREEN
    _set_east_asia_font(h3, '黑体')
    h3.paragraph_format.space_before   = Pt(8)
    h3.paragraph_format.space_after    = Pt(3)
    h3.paragraph_format.keep_with_next = True

    # ── Heading 4：深色，12pt，粗体 ────────────────────────────────────────
    h4 = _get_style(doc, 'Heading 4')
    h4.font.name      = 'Arial'
    h4.font.size      = Pt(12)
    h4.font.bold      = True
    h4.font.color.rgb = DARK_TEXT
    _set_east_asia_font(h4, '黑体')
    h4.paragraph_format.space_before   = Pt(6)
    h4.paragraph_format.space_after    = Pt(2)
    h4.paragraph_format.keep_with_next = True

    # ── Title：大标题 ─────────────────────────────────────────────────────
    title = _get_style(doc, 'Title')
    title.font.name      = 'Arial'
    title.font.size      = Pt(28)
    title.font.bold      = True
    title.font.color.rgb = BRAND_GREEN
    _set_east_asia_font(title, '黑体')

    # 占位段落
    doc.add_paragraph('BDAE ESA Report Template', style='Normal')

    doc.save(output_path)
    print(f'[OK] reference.docx saved → {output_path}')


def add_page_number(paragraph, color):
    """添加当前页码字段"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    run.font.size = Pt(8)
    run.font.color.rgb = color
    run.font.name = 'Arial'


def add_num_pages(paragraph, color):
    """添加总页数字段"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    run.font.size = Pt(8)
    run.font.color.rgb = color
    run.font.name = 'Arial'


def _get_style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, 1)


def _set_east_asia_font(style, font_name):
    rPr    = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'),       font_name)


if __name__ == '__main__':
    make_reference_docx()
