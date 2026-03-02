"""
════════════════════════════════════════════════════════════════════
ESA Report  MD → DOCX  Conversion Service
════════════════════════════════════════════════════════════════════
作者       BDAE Consulting
版本       2.0
最后更新   2026-03

技术路线：
  1. Dify 工作流生成完整 Markdown 字符串
  2. HTTP 请求节点 POST 到本服务 /convert
  3. Pandoc 将 Markdown 转换为 .docx（套用 reference.docx 模板样式）
  4. python-docx 后处理（动态内容：表头染色、HIGH/MED/LOW、列宽）
  5. 返回本地下载链接，Dify End 节点输出给用户

挂载目录：~/AI_Workspace/07_Pandoc-api  →  /app/templates
  - reference.docx  : Word 样式模板（可直接用 Word 编辑）
════════════════════════════════════════════════════════════════════
"""

from flask import Flask, request, jsonify, send_from_directory
import subprocess, tempfile, os, uuid, traceback

app = Flask(__name__)
OUTPUT_DIR     = '/tmp/pandoc_outputs'
REFERENCE_DOCX = '/app/templates/reference.docx'
os.makedirs(OUTPUT_DIR, exist_ok=True)


@app.route('/health', methods=['GET'])
def health():
    """供监控或 Dify 确认服务存活"""
    return jsonify({'status': 'ok'})


@app.route('/convert', methods=['POST'])
def convert():
    """
    接收 JSON，转换 Markdown → DOCX，返回下载链接。

    请求体（JSON）：
        markdown      : str  完整 Markdown 报告内容（必填）
        report_no     : str  报告编号，用于文件命名
        property_name : str  项目名称（预留，供后续页眉扩展）
        filename      : str  期望的文件名，默认 ESA_Report.docx

    返回（JSON）：
        success       : bool
        download_url  : str  本地可访问的下载链接
        filename      : str  实际生成的文件名
    """
    try:
        data          = request.get_json(force=True)
        md_content    = data.get('markdown', '')
        report_no     = data.get('report_no', '')
        property_name = data.get('property_name', '')
        filename      = data.get('filename', 'ESA_Report.docx')

        if not md_content:
            return jsonify({'success': False, 'error': 'markdown field is empty'}), 400

        # 生成唯一文件名，避免并发时文件覆盖
        file_id   = str(uuid.uuid4())[:8]
        base_name = filename.replace('.docx', '')
        safe_name = f"{base_name}_{file_id}.docx"
        docx_path = os.path.join(OUTPUT_DIR, safe_name)

        # ── Step 1: Pandoc 转换 MD → DOCX ───────────────────────────────
        with tempfile.NamedTemporaryFile(suffix='.md', mode='w',
                                         encoding='utf-8', delete=False) as f:
            f.write(md_content)
            md_path = f.name

        cmd = ['pandoc', md_path, '-o', docx_path, '--from=markdown', '--to=docx']
        if os.path.exists(REFERENCE_DOCX):
            cmd += [f'--reference-doc={REFERENCE_DOCX}']

        result = subprocess.run(cmd, capture_output=True, text=True)
        os.unlink(md_path)

        if result.returncode != 0:
            return jsonify({'success': False,
                            'error': f'Pandoc error: {result.stderr}'}), 500

        # ── Step 2: python-docx 后处理 ───────────────────────────────────
        postprocess(docx_path)

        return jsonify({
            'success':      True,
            'download_url': f'http://localhost:5050/files/{safe_name}',
            'filename':     safe_name
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e),
                        'trace': traceback.format_exc()}), 500


@app.route('/files/<filename>')
def download(filename):
    """提供转换后 .docx 文件的下载"""
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)


# ════════════════════════════════════════════════════════════════════
# 后处理
# 以下样式无法通过 reference.docx 模板实现，需要代码处理：
#   - 表头行背景色（Pandoc 不区分表头行）
#   - HIGH/MED/LOW 基于内容的动态染色
#   - 表格宽度和对齐方式
#   - 表格字体大小
# ════════════════════════════════════════════════════════════════════
def postprocess(docx_path):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(docx_path)

    # Risk Assessment 表格的 Ranking 列颜色映射
    RANK_COLORS = {
        'HIGH': ('C00000', 'FFFFFF'),  # 红底白字
        'MED':  ('FF8C00', 'FFFFFF'),  # 橙底白字
        'LOW':  ('375623', 'FFFFFF'),  # 深绿底白字
    }

    # A4 内容宽：(21cm - 2cm - 2cm) × 567 = 9639 DXA
    # 若 reference.docx 页边距不同，需同步修改此值
    PAGE_WIDTH_DXA = 9639

    for table in doc.tables:
        tbl   = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 表格宽度撑满页面内容宽
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'),    str(PAGE_WIDTH_DXA))
        tblW.set(qn('w:type'), 'dxa')

        # 表格左对齐（与正文标题对齐）
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'left')

        # 各列均分宽度
        col_count = len(table.columns)
        col_w     = PAGE_WIDTH_DXA // col_count if col_count > 0 else PAGE_WIDTH_DXA
        tblGrid   = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            for gc in tblGrid.findall(qn('w:gridCol')):
                gc.set(qn('w:w'), str(col_w))

        for ri, row in enumerate(table.rows):
            for cell in row.cells:

                # 同步单元格宽度
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW  = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'),    str(col_w))
                tcW.set(qn('w:type'), 'dxa')

                text = cell.text.strip().upper()
                if text in RANK_COLORS:
                    _set_cell_color(cell, *RANK_COLORS[text])
                elif ri == 0:
                    _set_cell_color(cell, '000000', 'FFFFFF')  # 表头黑底白字

                # 所有表格内容字体 10pt
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

    doc.save(docx_path)


def _set_cell_color(cell, bg_hex, fg_hex):
    """设置单元格背景色和文字颜色"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  bg_hex)
    tcPr.append(shd)

    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = _hex_to_rgb(fg_hex)
            run.bold = True


def _hex_to_rgb(hex_str):
    """十六进制颜色字符串 → RGBColor"""
    from docx.shared import RGBColor
    return RGBColor(int(hex_str[0:2], 16),
                    int(hex_str[2:4], 16),
                    int(hex_str[4:6], 16))


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5050, debug=False)
